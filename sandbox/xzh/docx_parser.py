import os
from typing import List, Dict, Optional
import json
from datetime import datetime

# Note: This module requires 'python-docx' to be installed.
# Please add 'python-docx' to your environment.yml and update the environment.
try:
    from docx import Document
except ImportError:
    print("Error: 'python-docx' module not found. Please install it via 'pip install python-docx' or add to environment.yml")
    Document = None

class DocxProcessor:
    """
    A robust processor for parsing .docx files, extracting text, metadata, and structure.
    Designed for the 'Service Outsourcing Competition' context.
    """

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.filename = os.path.basename(file_path)
        self.content: Dict = {"paragraphs": [], "tables": [], "metadata": {}}
        self._document = None

    def load(self) -> bool:
        """Loads the document from disk."""
        if not os.path.exists(self.file_path):
            print(f"File not found: {self.file_path}")
            return False
        
        if Document is None:
            return False

        try:
            self._document = Document(self.file_path)
            return True
        except Exception as e:
            print(f"Failed to load document: {e}")
            return False

    def parse(self) -> Dict:
        """
        Parses the document content into a structured format.
        """
        if not self._document:
            if not self.load():
                return {}

        # 1. Extract Core Properties
        core_props = self._document.core_properties
        self.content["metadata"] = {
            "author": core_props.author,
            "created": core_props.created.isoformat() if core_props.created else None,
            "modified": core_props.modified.isoformat() if core_props.modified else None,
            "title": core_props.title,
            "revision": core_props.revision
        }

        # 2. Extract Paragraphs with Style Info (Simulation of Structure)
        for i, para in enumerate(self._document.paragraphs):
            if para.text.strip():
                self.content["paragraphs"].append({
                    "id": i,
                    "text": para.text.strip(),
                    "style": para.style.name,
                    "is_heading": para.style.name.startswith('Heading')
                })

        # 3. Extract Tables
        for i, table in enumerate(self._document.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            self.content["tables"].append({
                "id": i,
                "data": table_data
            })

        return self.content

    def export_json(self, output_dir: str = "./output") -> str:
        """Exports the parsed content to JSON."""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        output_path = os.path.join(output_dir, f"{self.filename}.json")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(self.content, f, ensure_ascii=False, indent=2)
        return output_path

def main():
    # Target files specified by the user
    target_files = [
        r"f:\college\sophomore\服务外包\赛题信息\题目.docx",
        r"f:\college\sophomore\服务外包\赛题信息\老师的解析.docx",
        r"f:\college\sophomore\服务外包\赛题信息\分析.docx"
    ]

    processor_output_dir = r"f:\college\sophomore\服务外包\sandbox\xzh\parsed_docs"

    print("Starting Document Processing Task...")
    print("="*50)

    for file_path in target_files:
        print(f"Processing: {os.path.basename(file_path)}...")
        processor = DocxProcessor(file_path)
        
        result = processor.parse()
        if result:
            json_path = processor.export_json(processor_output_dir)
            print(f"  -> Success. Parsed {len(result['paragraphs'])} paragraphs.")
            print(f"  -> Exported to: {json_path}")
            
            # Print a brief preview (first 3 paragraphs)
            preview = [p['text'] for p in result['paragraphs'][:3]]
            print(f"  -> Preview: {preview}")
        else:
            print("  -> Failed to parse.")
        print("-" * 30)

if __name__ == "__main__":
    main()
